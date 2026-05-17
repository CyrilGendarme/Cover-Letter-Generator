"""Main Tkinter window and widgets."""

import tkinter as tk
from tkinter import ttk

from cover_letter_generator.models.text_parts import (
    TextPartKind,
    TextPartNonTech,
    Tech,
    TextPart,
    TextPartTech,
)
from cover_letter_generator.services.cover_letter_service import CoverLetterService
from cover_letter_generator.services.json_data_store import JsonDataStore


class MainWindow:
    """Builds and manages the main application window."""

    def __init__(self, service: CoverLetterService, data_store: JsonDataStore) -> None:
        self._service = service
        self._data_store = data_store
        self._root = tk.Tk()
        self._root.title("Cover Letter Generator")
        self._root.geometry("980x700")
        self._root.minsize(760, 560)

        self._position_name_var = tk.StringVar()
        self._position_company_var = tk.StringVar()
        self._status_var = tk.StringVar(value="Ready")
        self._include_english_var = tk.BooleanVar(value=True)
        self._include_french_var = tk.BooleanVar(value=False)

        # Export directory - default to Downloads
        from pathlib import Path

        default_export_dir = str(Path.home() / "Downloads")
        self._export_directory_var = tk.StringVar(value=default_export_dir)

        self._tech_name_var = tk.StringVar()
        self._techs: list[Tech] = []
        self._tech_drag_start_index: int | None = None

        self._text_parts: list[TextPart] = []
        self._text_part_name_var = tk.StringVar()
        self._text_part_association_var = tk.StringVar(value="none")
        self._text_part_enum_var = tk.StringVar(value=TextPartKind.INTRO.value)
        self._text_part_tech_var = tk.StringVar(value="")
        self._text_part_always_include_var = tk.BooleanVar(value=False)
        self._selected_tech_part_vars: dict[int, tk.BooleanVar] = {}
        self._generated_output_with_tags = ""
        self._application_reason_items: list[tuple[int, TextPartNonTech]] = []

        self._screens: ttk.Notebook | None = None

        self._build_layout()
        self._load_configuration()

    def _build_layout(self) -> None:
        main = ttk.Frame(self._root, padding=16)
        main.pack(fill=tk.BOTH, expand=True)

        title = ttk.Label(main, text="Cover Letter Generator", font=("Segoe UI", 20, "bold"))
        title.pack(anchor=tk.W)

        subtitle = ttk.Label(
            main,
            text="Fill in the form and generate a first draft.",
            font=("Segoe UI", 10),
        )
        subtitle.pack(anchor=tk.W, pady=(0, 12))

        notebook = ttk.Notebook(main)
        notebook.pack(fill=tk.BOTH, expand=True)

        generator_tab = ttk.Frame(notebook, padding=8)
        configuration_tab = ttk.Frame(notebook, padding=8)

        notebook.add(generator_tab, text="Generator")
        notebook.add(configuration_tab, text="Configuration")

        self._build_generator_tab(generator_tab)
        self._build_configuration_tab(configuration_tab)

        status = ttk.Label(main, textvariable=self._status_var)
        status.pack(anchor=tk.W, pady=(8, 0))

    def _build_generator_tab(self, parent: ttk.Frame) -> None:
        self._screens = ttk.Notebook(parent)
        self._screens.pack(fill=tk.BOTH, expand=True)

        screen_1 = ttk.Frame(self._screens, padding=8)
        screen_2 = ttk.Frame(self._screens, padding=8)
        screen_3 = ttk.Frame(self._screens, padding=8)

        self._screens.add(screen_1, text="1) Job Description")
        self._screens.add(screen_2, text="2) Generation Setup")
        self._screens.add(screen_3, text="3) Output")

        self._build_screen_job_description(screen_1)
        self._build_screen_generation_setup(screen_2)
        self._build_screen_output(screen_3)
        self._screens.select(0)

    def _build_screen_job_description(self, parent: ttk.Frame) -> None:
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(1, weight=1)

        ttk.Label(parent, text="Paste job description").grid(row=0, column=0, sticky=tk.W)
        self._job_description = tk.Text(parent, height=12, wrap=tk.WORD)
        self._job_description.grid(row=1, column=0, sticky=tk.NSEW, pady=(4, 8))

        actions = ttk.Frame(parent)
        actions.grid(row=2, column=0, sticky=tk.W, pady=(0, 8))
        ttk.Button(actions, text="Prepare matches", command=self._prepare_text_part_matches).pack(side=tk.LEFT)
        ttk.Button(actions, text="Clear", command=self._clear_job_description).pack(side=tk.LEFT, padx=(6, 0))

    def _build_screen_generation_setup(self, parent: ttk.Frame) -> None:
        parent.columnconfigure(1, weight=1)
        parent.rowconfigure(3, weight=1)
        parent.rowconfigure(4, weight=1)

        self._add_labeled_entry(parent, "Job title", self._position_name_var, row=0)
        self._add_labeled_entry(parent, "Company name", self._position_company_var, row=1)

        ttk.Label(parent, text="Output language(s)").grid(
            row=2, column=0, sticky=tk.W, padx=(0, 8), pady=4
        )
        language_frame = ttk.Frame(parent)
        language_frame.grid(row=2, column=1, sticky=tk.W, pady=4)
        ttk.Checkbutton(
            language_frame,
            text="English",
            variable=self._include_english_var,
        ).pack(side=tk.LEFT)
        ttk.Checkbutton(
            language_frame,
            text="French",
            variable=self._include_french_var,
        ).pack(side=tk.LEFT, padx=(12, 0))

        ttk.Label(parent, text="TextPartTech items").grid(
            row=3, column=0, sticky=tk.NW, padx=(0, 8), pady=4
        )
        tech_parts_frame = ttk.LabelFrame(parent, text="Check the parts to include")
        tech_parts_frame.grid(row=3, column=1, sticky=tk.NSEW, pady=4)
        tech_parts_frame.columnconfigure(0, weight=1)
        tech_parts_frame.rowconfigure(0, weight=1)

        self._tech_parts_canvas = tk.Canvas(tech_parts_frame, borderwidth=0, highlightthickness=0)
        self._tech_parts_canvas.grid(row=0, column=0, sticky=tk.NSEW)

        tech_parts_scrollbar = ttk.Scrollbar(tech_parts_frame, orient=tk.VERTICAL, command=self._tech_parts_canvas.yview)
        tech_parts_scrollbar.grid(row=0, column=1, sticky=tk.NS)
        self._tech_parts_canvas.configure(yscrollcommand=tech_parts_scrollbar.set)

        self._tech_parts_check_frame = ttk.Frame(self._tech_parts_canvas)
        self._tech_parts_canvas_window = self._tech_parts_canvas.create_window((0, 0), window=self._tech_parts_check_frame, anchor="nw")
        self._tech_parts_check_frame.bind("<Configure>", self._on_tech_parts_frame_configure)
        self._tech_parts_canvas.bind("<Configure>", self._on_tech_parts_canvas_configure)

        ttk.Label(parent, text="Application reason").grid(
            row=4, column=0, sticky=tk.NW, padx=(0, 8), pady=4
        )
        application_reason_frame = ttk.LabelFrame(
            parent, text="Select one application reason"
        )
        application_reason_frame.grid(row=4, column=1, sticky=tk.NSEW, pady=4)
        application_reason_frame.columnconfigure(0, weight=1)
        application_reason_frame.rowconfigure(0, weight=1)

        self._application_reason_list = tk.Listbox(
            application_reason_frame, exportselection=False
        )
        self._application_reason_list.grid(
            row=0, column=0, sticky=tk.NSEW, padx=(6, 0), pady=6
        )

        application_reason_scrollbar = ttk.Scrollbar(
            application_reason_frame,
            orient=tk.VERTICAL,
            command=self._application_reason_list.yview,
        )
        application_reason_scrollbar.grid(
            row=0, column=1, sticky=tk.NS, padx=(0, 6), pady=6
        )
        self._application_reason_list.configure(
            yscrollcommand=application_reason_scrollbar.set
        )

        actions = ttk.Frame(parent)
        actions.grid(row=5, column=1, sticky=tk.W, pady=(8, 0))
        ttk.Button(actions, text="Refresh matches", command=self._prepare_text_part_matches).pack(side=tk.LEFT)
        ttk.Button(actions, text="Generate text", command=self._on_generate).pack(side=tk.LEFT, padx=(6, 0))

    def _build_screen_output(self, parent: ttk.Frame) -> None:
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(0, weight=1)

        self._output = tk.Text(parent, wrap=tk.WORD)
        self._output.grid(row=0, column=0, sticky=tk.NSEW)

        actions = ttk.Frame(parent)
        actions.grid(row=1, column=0, sticky=tk.W, pady=(8, 0))
        ttk.Button(actions, text="Copy", command=self._copy_output).pack(side=tk.LEFT)
        ttk.Button(actions, text="Export Word", command=self._export_word).pack(side=tk.LEFT, padx=(6, 0))
        ttk.Button(actions, text="Export PDF", command=self._export_pdf).pack(side=tk.LEFT, padx=(6, 0))

    def _build_configuration_tab(self, parent: ttk.Frame) -> None:
        parent.columnconfigure(0, weight=1)
        parent.columnconfigure(1, weight=1)
        parent.rowconfigure(0, weight=0)
        parent.rowconfigure(1, weight=1)

        # Export directory section
        export_frame = ttk.LabelFrame(parent, text="Export Settings")
        export_frame.grid(
            row=0, column=0, columnspan=2, sticky=tk.EW, padx=6, pady=(6, 12)
        )
        export_frame.columnconfigure(1, weight=1)

        ttk.Label(export_frame, text="Export directory:").grid(
            row=0, column=0, sticky=tk.W, padx=8, pady=8
        )
        ttk.Entry(export_frame, textvariable=self._export_directory_var).grid(
            row=0, column=1, sticky=tk.EW, padx=(0, 6), pady=8
        )
        ttk.Button(
            export_frame, text="Browse", command=self._browse_export_directory
        ).grid(row=0, column=2, padx=(0, 6), pady=8)
        ttk.Button(export_frame, text="Open", command=self._open_export_directory).grid(
            row=0, column=3, padx=(0, 8), pady=8
        )

        tech_frame = ttk.LabelFrame(parent, text="Tech Entities")
        tech_frame.grid(row=1, column=0, sticky=tk.NSEW, padx=(0, 6))

        text_part_frame = ttk.LabelFrame(parent, text="TextPart Entities")
        text_part_frame.grid(row=1, column=1, sticky=tk.NSEW, padx=(6, 0))

        self._build_tech_editor(tech_frame)
        self._build_text_part_editor(text_part_frame)

    def _build_tech_editor(self, parent: ttk.LabelFrame) -> None:
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(0, weight=1)

        self._tech_list = tk.Listbox(parent, exportselection=False)
        self._tech_list.grid(row=0, column=0, columnspan=2, sticky=tk.NSEW, padx=8, pady=(8, 4))
        self._tech_list.bind("<<ListboxSelect>>", self._on_tech_selected)
        self._tech_list.bind("<ButtonPress-1>", self._on_tech_drag_start)
        self._tech_list.bind("<B1-Motion>", self._on_tech_drag_motion)
        self._tech_list.bind("<ButtonRelease-1>", self._on_tech_drag_release)

        ttk.Label(parent, text="Name").grid(row=1, column=0, sticky=tk.W, padx=8)
        ttk.Entry(parent, textvariable=self._tech_name_var).grid(row=2, column=0, columnspan=2, sticky=tk.EW, padx=8)

        ttk.Label(parent, text="Associated text parts (one per line)").grid(row=3, column=0, sticky=tk.W, padx=8, pady=(8, 0))
        self._tech_text_parts = tk.Text(parent, height=8, wrap=tk.WORD)
        self._tech_text_parts.grid(row=4, column=0, columnspan=2, sticky=tk.EW, padx=8, pady=(2, 8))

        buttons = ttk.Frame(parent)
        buttons.grid(row=5, column=0, columnspan=2, sticky=tk.W, padx=8, pady=(0, 8))
        ttk.Button(buttons, text="Load", command=self._load_configuration).pack(side=tk.LEFT)
        ttk.Button(buttons, text="Add / Modify", command=self._on_upsert_tech).pack(side=tk.LEFT, padx=(6, 0))
        ttk.Button(buttons, text="Remove", command=self._on_remove_tech).pack(side=tk.LEFT, padx=(6, 0))

    def _build_text_part_editor(self, parent: ttk.LabelFrame) -> None:
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(0, weight=1)

        self._text_part_list = tk.Listbox(parent, exportselection=False)
        self._text_part_list.grid(row=0, column=0, sticky=tk.NSEW, padx=8, pady=(8, 4))
        self._text_part_list.bind("<<ListboxSelect>>", self._on_text_part_selected)

        ttk.Label(parent, text="Name (optional)").grid(
            row=1, column=0, sticky=tk.W, padx=8
        )
        ttk.Entry(parent, textvariable=self._text_part_name_var).grid(
            row=2, column=0, sticky=tk.EW, padx=8
        )

        ttk.Label(parent, text="Texts (English / French)").grid(
            row=3, column=0, sticky=tk.W, padx=8
        )
        text_frame = ttk.Frame(parent)
        text_frame.grid(row=4, column=0, sticky=tk.NSEW, padx=8, pady=(2, 8))
        text_frame.columnconfigure(0, weight=1)
        text_frame.columnconfigure(1, weight=1)

        ttk.Label(text_frame, text="English").grid(row=0, column=0, sticky=tk.W)
        ttk.Label(text_frame, text="French").grid(row=0, column=1, sticky=tk.W)

        self._text_part_text = tk.Text(text_frame, height=12, wrap=tk.WORD)
        self._text_part_text.grid(row=1, column=0, sticky=tk.NSEW, padx=(0, 6))
        self._text_part_french_text = tk.Text(text_frame, height=12, wrap=tk.WORD)
        self._text_part_french_text.grid(row=1, column=1, sticky=tk.NSEW, padx=(6, 0))

        ttk.Label(parent, text="Association type").grid(
            row=5, column=0, sticky=tk.W, padx=8
        )
        association_frame = ttk.Frame(parent)
        association_frame.grid(row=6, column=0, sticky=tk.W, padx=8, pady=(2, 6))
        ttk.Radiobutton(
            association_frame,
            text="None",
            value="none",
            variable=self._text_part_association_var,
            command=self._toggle_text_part_association_inputs,
        ).pack(side=tk.LEFT)
        ttk.Radiobutton(
            association_frame,
            text="Tech",
            value="tech",
            variable=self._text_part_association_var,
            command=self._toggle_text_part_association_inputs,
        ).pack(side=tk.LEFT, padx=(8, 0))
        ttk.Radiobutton(
            association_frame,
            text="Enum",
            value="enum",
            variable=self._text_part_association_var,
            command=self._toggle_text_part_association_inputs,
        ).pack(side=tk.LEFT, padx=(8, 0))

        enum_row = ttk.Frame(parent)
        enum_row.grid(row=7, column=0, sticky=tk.W, padx=8, pady=(0, 4))
        ttk.Label(enum_row, text="Enum value").pack(side=tk.LEFT)
        self._text_part_enum_combo = ttk.Combobox(
            enum_row,
            state="readonly",
            values=[item.value for item in TextPartKind],
            textvariable=self._text_part_enum_var,
            width=12,
        )
        self._text_part_enum_combo.pack(side=tk.LEFT, padx=(8, 0))

        tech_row = ttk.Frame(parent)
        tech_row.grid(row=8, column=0, sticky=tk.W, padx=8, pady=(0, 8))
        ttk.Label(tech_row, text="Tech key").pack(side=tk.LEFT)
        self._text_part_tech_combo = ttk.Combobox(
            tech_row,
            state="readonly",
            textvariable=self._text_part_tech_var,
            width=20,
        )
        self._text_part_tech_combo.pack(side=tk.LEFT, padx=(8, 0))

        self._text_part_always_include_check = ttk.Checkbutton(
            parent,
            text="Always include in letter (hidden from checklist)",
            variable=self._text_part_always_include_var,
        )
        self._text_part_always_include_check.grid(
            row=9, column=0, sticky=tk.W, padx=8, pady=(0, 8)
        )

        buttons = ttk.Frame(parent)
        buttons.grid(row=10, column=0, sticky=tk.W, padx=8, pady=(0, 8))
        ttk.Button(buttons, text="Load", command=self._load_configuration).pack(side=tk.LEFT)
        ttk.Button(buttons, text="Add / Modify", command=self._on_upsert_text_part).pack(side=tk.LEFT, padx=(6, 0))
        ttk.Button(buttons, text="Remove", command=self._on_remove_text_part).pack(side=tk.LEFT, padx=(6, 0))
        self._toggle_text_part_association_inputs()

    def _load_configuration(self) -> None:
        self._techs = self._data_store.load_techs()
        self._text_parts = self._data_store.load_text_parts()
        saved_export_directory = self._data_store.load_export_directory()
        if saved_export_directory:
            self._export_directory_var.set(saved_export_directory)
        self._refresh_tech_list()
        self._refresh_text_part_list()
        self._refresh_text_part_tech_keys()
        self._refresh_text_part_tech_candidates()
        self._refresh_application_reason_candidates()
        self._status_var.set("Configuration loaded from JSON files")

    def _refresh_text_part_tech_keys(self) -> None:
        tech_names = [tech.name for tech in self._techs if tech.name.strip()]
        self._text_part_tech_combo["values"] = tech_names
        if tech_names and self._text_part_tech_var.get() not in tech_names:
            self._text_part_tech_var.set(tech_names[0])
        if not tech_names:
            self._text_part_tech_var.set("")

    def _refresh_text_part_tech_candidates(self) -> None:
        tech_parts = self._get_checkable_text_part_tech_items()
        current_states = {
            index: variable.get() for index, variable in self._selected_tech_part_vars.items()
        }
        self._selected_tech_part_vars = {
            index: tk.BooleanVar(value=current_states.get(index, False)) for index, _part in tech_parts
        }
        self._render_text_part_tech_checkboxes(tech_parts)

    def _refresh_tech_list(self) -> None:
        self._tech_list.delete(0, tk.END)
        for tech in self._techs:
            self._tech_list.insert(tk.END, f"{tech.name} ({len(tech.text_parts)} text parts)")

    def _refresh_text_part_list(self) -> None:
        self._text_part_list.delete(0, tk.END)
        for part in self._text_parts:
            preview = part.text.replace("\n", " ").strip()
            if len(preview) > 70:
                preview = f"{preview[:67]}..."
            label = self._format_text_part_association(part)
            identifier = self._text_part_identifier(part)
            self._text_part_list.insert(
                tk.END, f"[{label}] {identifier} -> {preview or '<empty>'}"
            )

    def _get_text_part_tech_items(self) -> list[tuple[int, TextPartTech]]:
        items: list[tuple[int, TextPartTech]] = [
            (index, part)
            for index, part in enumerate(self._text_parts)
            if isinstance(part, TextPartTech)
        ]

        tech_order = {
            tech.name.strip().lower(): order for order, tech in enumerate(self._techs)
        }

        items.sort(
            key=lambda item: (
                tech_order.get(item[1].tech.name.strip().lower(), len(self._techs)),
                item[0],
            )
        )
        return items

    def _get_checkable_text_part_tech_items(self) -> list[tuple[int, TextPartTech]]:
        return [
            (index, part)
            for index, part in self._get_text_part_tech_items()
            if not part.always_include
        ]

    def _get_all_non_tech_parts(self) -> list[TextPartNonTech]:
        return [part for part in self._text_parts if isinstance(part, TextPartNonTech)]

    def _get_application_reason_items(self) -> list[tuple[int, TextPartNonTech]]:
        return [
            (index, part)
            for index, part in enumerate(self._text_parts)
            if isinstance(part, TextPartNonTech)
            and part.text_part_kind == TextPartKind.APPLICATION_REASON
        ]

    def _refresh_application_reason_candidates(self) -> None:
        selected_part: TextPartNonTech | None = None
        current_selection = self._application_reason_list.curselection()
        if current_selection and self._application_reason_items:
            selected_list_index = int(current_selection[0])
            if 0 <= selected_list_index < len(self._application_reason_items):
                _, selected_part = self._application_reason_items[selected_list_index]

        self._application_reason_items = self._get_application_reason_items()
        self._application_reason_list.delete(0, tk.END)

        for _, part in self._application_reason_items:
            identifier = self._text_part_identifier(part)
            label = f"{identifier}: {self._truncate_text(part.text, limit=120)}"
            self._application_reason_list.insert(tk.END, label)

        if not self._application_reason_items:
            self._application_reason_list.insert(
                tk.END, "No application reason text parts configured yet."
            )
            self._application_reason_list.configure(state=tk.DISABLED)
            return

        self._application_reason_list.configure(state=tk.NORMAL)

        selected_index = 0
        if selected_part is not None:
            for index, (_, part) in enumerate(self._application_reason_items):
                if part.text == selected_part.text and part.name == selected_part.name:
                    selected_index = index
                    break

        self._application_reason_list.selection_clear(0, tk.END)
        self._application_reason_list.selection_set(selected_index)
        self._application_reason_list.activate(selected_index)
        self._application_reason_list.see(selected_index)

    def _get_selected_application_reason_part(self) -> TextPartNonTech | None:
        if not self._application_reason_items:
            return None
        selected = self._application_reason_list.curselection()
        selected_index = int(selected[0]) if selected else 0
        if selected_index < 0 or selected_index >= len(self._application_reason_items):
            selected_index = 0
        _, part = self._application_reason_items[selected_index]
        return part

    def _get_non_tech_parts_for_generation(
        self,
        exclude_eager_to_relocate: bool = False,
    ) -> list[TextPartNonTech]:
        selected_application_reason = self._get_selected_application_reason_part()
        result: list[TextPartNonTech] = []
        for part in self._get_all_non_tech_parts():
            if part.text_part_kind == TextPartKind.APPLICATION_REASON:
                continue
            if (
                exclude_eager_to_relocate
                and part.text_part_kind == TextPartKind.EAGER_TO_RELOCATE
            ):
                continue
            result.append(part)

        if selected_application_reason is not None:
            result.append(selected_application_reason)

        return result

    def _render_text_part_tech_checkboxes(self, tech_parts: list[tuple[int, TextPartTech]]) -> None:
        for child in self._tech_parts_check_frame.winfo_children():
            child.destroy()

        if not tech_parts:
            ttk.Label(
                self._tech_parts_check_frame,
                text="No TextPartTech items configured yet.",
            ).pack(anchor=tk.W, padx=6, pady=6)
            return

        for index, part in tech_parts:
            identifier = self._text_part_identifier(part)
            label = f"{identifier}: {self._truncate_text(part.text)}"
            ttk.Checkbutton(
                self._tech_parts_check_frame,
                text=label,
                variable=self._selected_tech_part_vars[index],
            ).pack(anchor=tk.W, padx=6, pady=2)

    def _truncate_text(self, text: str, limit: int = 80) -> str:
        cleaned = text.replace("\n", " ").strip()
        if len(cleaned) <= limit:
            return cleaned
        return f"{cleaned[: limit - 3]}..."

    def _format_text_part_association(self, part: TextPart) -> str:
        if isinstance(part, TextPartTech):
            mode = "always" if part.always_include else "checkable"
            return f"tech:{part.tech.name}:{mode}"
        if isinstance(part, TextPartNonTech):
            return f"enum:{part.text_part_kind.value}"
        return "none"

    def _text_part_identifier(self, part: TextPart) -> str:
        return part.identifier()

    def _on_tech_selected(self, _event: tk.Event) -> None:
        index = self._get_selected_index(self._tech_list)
        if index is None:
            return
        selected = self._techs[index]
        self._tech_name_var.set(selected.name)
        self._tech_text_parts.delete("1.0", tk.END)
        self._tech_text_parts.insert("1.0", "\n".join(selected.text_parts))

    def _on_tech_drag_start(self, event: tk.Event) -> None:
        if self._tech_list.size() == 0:
            self._tech_drag_start_index = None
            return
        index = self._tech_list.nearest(event.y)
        if index < 0:
            self._tech_drag_start_index = None
            return
        self._tech_drag_start_index = index

    def _on_tech_drag_motion(self, event: tk.Event) -> None:
        if self._tech_drag_start_index is None or self._tech_list.size() == 0:
            return
        target_index = self._tech_list.nearest(event.y)
        if target_index < 0:
            return
        self._tech_list.selection_clear(0, tk.END)
        self._tech_list.selection_set(target_index)
        self._tech_list.activate(target_index)

    def _on_tech_drag_release(self, event: tk.Event) -> None:
        if self._tech_drag_start_index is None or self._tech_list.size() == 0:
            self._tech_drag_start_index = None
            return

        source_index = self._tech_drag_start_index
        self._tech_drag_start_index = None
        target_index = self._tech_list.nearest(event.y)
        if target_index < 0 or target_index >= len(self._techs):
            return
        if source_index == target_index:
            return

        self._move_tech(source_index, target_index)

    def _move_tech(self, source_index: int, target_index: int) -> None:
        moved = self._techs.pop(source_index)
        self._techs.insert(target_index, moved)
        self._data_store.save_techs(self._techs)

        self._refresh_tech_list()
        self._tech_list.selection_clear(0, tk.END)
        self._tech_list.selection_set(target_index)
        self._tech_list.activate(target_index)
        self._tech_list.see(target_index)

        self._refresh_text_part_tech_keys()
        self._refresh_text_part_tech_candidates()
        self._status_var.set("Tech order updated")

    def _on_text_part_selected(self, _event: tk.Event) -> None:
        index = self._get_selected_index(self._text_part_list)
        if index is None:
            return
        selected = self._text_parts[index]
        self._text_part_name_var.set(selected.name)
        self._text_part_text.delete("1.0", tk.END)
        self._text_part_text.insert("1.0", selected.text)
        self._text_part_french_text.delete("1.0", tk.END)
        self._text_part_french_text.insert("1.0", selected.french_text)

        if isinstance(selected, TextPartTech):
            self._text_part_association_var.set("tech")
            self._text_part_tech_var.set(selected.tech.name)
            self._text_part_always_include_var.set(selected.always_include)
        elif isinstance(selected, TextPartNonTech):
            self._text_part_association_var.set("enum")
            self._text_part_enum_var.set(selected.text_part_kind.value)
            self._text_part_always_include_var.set(False)
        else:
            self._text_part_association_var.set("none")
            self._text_part_always_include_var.set(False)
        self._toggle_text_part_association_inputs()

    def _on_upsert_tech(self) -> None:
        name = self._tech_name_var.get().strip()
        if not name:
            self._status_var.set("Tech name is required")
            return
        text_parts = self._read_non_empty_lines(self._tech_text_parts)
        result = self._data_store.upsert_tech_by_name(Tech(name=name, text_parts=text_parts))
        self._load_configuration()
        self._status_var.set(f"{result.capitalize()} tech by key '{name}'")

    def _on_remove_tech(self) -> None:
        index = self._get_selected_index(self._tech_list)
        if index is None:
            self._status_var.set("Select a tech entity to remove")
            return
        success = self._data_store.remove_tech(index)
        if success:
            self._load_configuration()
            self._tech_name_var.set("")
            self._tech_text_parts.delete("1.0", tk.END)
            self._status_var.set("Removed tech entity")
        else:
            self._status_var.set("Could not remove selected tech")

    def _on_upsert_text_part(self) -> None:
        name = self._text_part_name_var.get().strip()
        # Use end-1c to drop Tkinter's implicit trailing newline while preserving user spacing.
        text = self._text_part_text.get("1.0", "end-1c")
        french_text = self._text_part_french_text.get("1.0", "end-1c")
        if not text.strip():
            self._status_var.set("TextPart text is required")
            return

        association = self._text_part_association_var.get()
        if association == "tech":
            tech_key = self._text_part_tech_var.get().strip()
            if not tech_key:
                self._status_var.set("Select a tech key for TextPartTech association")
                return
            text_part = TextPartTech(
                text=text,
                name=name,
                french_text=french_text,
                tech=Tech(name=tech_key),
                always_include=self._text_part_always_include_var.get(),
            )
        elif association == "enum":
            enum_value = self._text_part_enum_var.get().strip().lower()
            if enum_value not in {item.value for item in TextPartKind}:
                self._status_var.set("Select a valid enum value")
                return
            text_part = TextPartNonTech(
                text=text,
                name=name,
                french_text=french_text,
                text_part_kind=TextPartKind(enum_value),
            )
        else:
            text_part = TextPart(text=text, name=name, french_text=french_text)

        result = self._data_store.upsert_text_part_by_text(text_part)
        self._load_configuration()
        self._status_var.set(f"{result.capitalize()} text part by key")

    def _toggle_text_part_association_inputs(self) -> None:
        association = self._text_part_association_var.get()
        if association == "enum":
            self._text_part_enum_combo.configure(state="readonly")
            self._text_part_tech_combo.configure(state="disabled")
            self._text_part_always_include_check.configure(state="disabled")
            self._text_part_always_include_var.set(False)
        elif association == "tech":
            self._text_part_enum_combo.configure(state="disabled")
            self._text_part_tech_combo.configure(state="readonly")
            self._text_part_always_include_check.configure(state="normal")
        else:
            self._text_part_enum_combo.configure(state="disabled")
            self._text_part_tech_combo.configure(state="disabled")
            self._text_part_always_include_check.configure(state="disabled")
            self._text_part_always_include_var.set(False)

    def _on_remove_text_part(self) -> None:
        index = self._get_selected_index(self._text_part_list)
        if index is None:
            self._status_var.set("Select a text part entity to remove")
            return
        success = self._data_store.remove_text_part(index)
        if success:
            self._load_configuration()
            self._text_part_name_var.set("")
            self._text_part_text.delete("1.0", tk.END)
            self._text_part_french_text.delete("1.0", tk.END)
            self._status_var.set("Removed text part")
        else:
            self._status_var.set("Could not remove selected text part")

    def _read_non_empty_lines(self, widget: tk.Text) -> list[str]:
        return [line.strip() for line in widget.get("1.0", tk.END).splitlines() if line.strip()]

    def _prepare_text_part_matches(self) -> None:
        job_description = self._job_description.get("1.0", tk.END).strip().lower()
        tech_parts = self._get_checkable_text_part_tech_items()
        match_count = 0
        for index, part in tech_parts:
            should_check = self._matches_job_description(job_description, part)
            self._selected_tech_part_vars[index].set(should_check)
            if should_check:
                match_count += 1
        self._status_var.set(f"Prepared {match_count} matching TextPartTech items")

        if self._screens:
            self._screens.select(1)

    def _matches_job_description(self, job_description: str, part: TextPartTech) -> bool:
        if not job_description:
            return False

        tech_name = part.tech.name.strip().lower()
        if tech_name and tech_name in job_description:
            return True

        for tech in self._techs:
            if tech.name.strip().lower() != tech_name:
                continue
            for alias in tech.text_parts:
                if alias.strip().lower() in job_description:
                    return True
        return False

    def _on_tech_parts_frame_configure(self, _event: tk.Event) -> None:
        self._tech_parts_canvas.configure(scrollregion=self._tech_parts_canvas.bbox("all"))

    def _on_tech_parts_canvas_configure(self, event: tk.Event) -> None:
        self._tech_parts_canvas.itemconfigure(self._tech_parts_canvas_window, width=event.width)

    def _clear_job_description(self) -> None:
        self._job_description.delete("1.0", tk.END)
        for variable in self._selected_tech_part_vars.values():
            variable.set(False)
        self._status_var.set("Job description cleared")

    def _get_selected_index(self, listbox: tk.Listbox) -> int | None:
        selected = listbox.curselection()
        if not selected:
            return None
        return int(selected[0])

    def _add_labeled_entry(
        self,
        parent: ttk.Frame,
        label: str,
        variable: tk.StringVar,
        row: int,
    ) -> None:
        ttk.Label(parent, text=label).grid(row=row, column=0, sticky=tk.W, padx=(0, 8), pady=4)
        ttk.Entry(parent, textvariable=variable).grid(row=row, column=1, sticky=tk.EW, pady=4)

    def _on_generate(self) -> None:
        include_english = self._include_english_var.get()
        include_french = self._include_french_var.get()
        if not include_english and not include_french:
            self._status_var.set("Select at least one output language")
            return

        is_french_only = include_french and not include_english

        selected_tech_parts: list[TextPartTech] = []
        for index, part in self._get_text_part_tech_items():
            if part.always_include:
                selected_tech_parts.append(part)
                continue
            variable = self._selected_tech_part_vars.get(index)
            if variable is not None and variable.get():
                selected_tech_parts.append(part)

        generated_outputs: list[str] = []
        if include_english:
            english_text = self._service.compose_cover_letter(
                job_title=self._position_name_var.get().strip(),
                company=self._position_company_var.get().strip(),
                tech_parts=selected_tech_parts,
                non_tech_parts=self._get_non_tech_parts_for_generation(),
                language="english",
            )
            generated_outputs.append(english_text)

        if include_french:
            french_text = self._service.compose_cover_letter(
                job_title=self._position_name_var.get().strip(),
                company=self._position_company_var.get().strip(),
                tech_parts=selected_tech_parts,
                non_tech_parts=self._get_non_tech_parts_for_generation(
                    exclude_eager_to_relocate=is_french_only
                ),
                language="french",
            )
            generated_outputs.append(french_text)

        text = "\n\n".join(generated_outputs)
        self._generated_output_with_tags = text
        display_text = self._strip_bold_tags(text)
        self._output.delete("1.0", tk.END)
        self._output.insert("1.0", display_text)
        self._status_var.set("Cover letter generated")

        if self._screens:
            self._screens.select(2)

    def _on_clear(self) -> None:
        self._position_name_var.set("")
        self._position_company_var.set("")
        self._include_english_var.set(True)
        self._include_french_var.set(False)
        self._job_description.delete("1.0", tk.END)
        for variable in self._selected_tech_part_vars.values():
            variable.set(False)
        self._generated_output_with_tags = ""
        self._output.delete("1.0", tk.END)
        self._status_var.set("Generator form cleared")

    def _browse_export_directory(self) -> None:
        """Allow user to browse and select export directory."""
        from tkinter import filedialog
        from pathlib import Path

        current_dir = self._export_directory_var.get()
        selected_dir = filedialog.askdirectory(
            title="Select Export Directory",
            initialdir=current_dir if Path(current_dir).exists() else str(Path.home()),
        )
        if selected_dir:
            self._export_directory_var.set(selected_dir)
            self._data_store.save_export_directory(selected_dir)
            self._status_var.set(f"Export directory set to: {selected_dir}")

    def _open_export_directory(self) -> None:
        """Open export directory in Windows Explorer."""
        import subprocess
        from pathlib import Path

        export_dir = self._export_directory_var.get()
        if not export_dir:
            self._status_var.set("No export directory selected")
            return

        path = Path(export_dir)
        if not path.exists():
            path.mkdir(parents=True, exist_ok=True)

        try:
            subprocess.Popen(f'explorer "{export_dir}"')
            self._status_var.set(f"Opened: {export_dir}")
        except Exception as e:
            self._status_var.set(f"Error opening directory: {e}")

    def _copy_output(self) -> None:
        text = self._output.get("1.0", tk.END).strip()
        if not text:
            self._status_var.set("No generated text to copy")
            return
        self._root.clipboard_clear()
        self._root.clipboard_append(text)
        self._status_var.set("Generated text copied to clipboard")

    def _export_word(self) -> None:
        self._status_var.set("Exporting Word...")
        self._root.update()  # Force GUI update
        try:
            text = self._get_export_source_text()
            if not text:
                self._status_var.set("No generated text to export")
                return

            job_title = self._position_name_var.get().strip()
            if not job_title:
                self._status_var.set("Job title is required to export")
                return

            from docx import Document
            from docx.shared import Pt
            from pathlib import Path
            import re

            safe_title = re.sub(r'[<>:"/\\|?*]', "", job_title)
            filename = f"Cover letter - {safe_title} - Cyril Gendarme.docx"
            export_dir = Path(self._export_directory_var.get())
            export_dir.mkdir(parents=True, exist_ok=True)
            self._data_store.save_export_directory(str(export_dir))
            path = export_dir / filename

            doc = Document()
            normal_style = doc.styles["Normal"]
            normal_style.font.name = "Calibri"
            normal_style.font.size = Pt(11)
            normal_style.paragraph_format.space_before = Pt(0)
            normal_style.paragraph_format.space_after = Pt(0)
            normal_style.paragraph_format.line_spacing = 1
            import re as regex_module

            for line in text.split("\n"):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1

                # Preserve empty lines as spacing paragraphs.
                if line == "":
                    continue

                # Check if this is a tech bullet (- Name: content)
                match = regex_module.match(r"^(\s*-\s*)([^:]+)(:)(.*)$", line)
                if match:
                    prefix = match.group(1)
                    name = match.group(2)
                    colon = match.group(3)
                    content = match.group(4)

                    p.add_run(prefix)
                    self._add_word_runs_with_bold_tags(p, name, default_bold=True)
                    self._add_word_runs_with_bold_tags(p, colon + content)
                else:
                    self._add_word_runs_with_bold_tags(p, line)

            doc.save(str(path))
            self._status_var.set(f"Word export complete: {path}")
        except ImportError as e:
            self._status_var.set("python-docx is not installed")
        except Exception as e:
            self._status_var.set(f"Export error: {e}")

    def _export_pdf(self) -> None:
        self._status_var.set("Exporting PDF...")
        self._root.update()  # Force GUI update
        try:
            text = self._get_export_source_text()
            if not text:
                self._status_var.set("No generated text to export")
                return

            job_title = self._position_name_var.get().strip()
            if not job_title:
                self._status_var.set("Job title is required to export")
                return

            from reportlab.lib.pagesizes import A4
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            from pathlib import Path
            import re

            safe_title = re.sub(r'[<>:"/\\|?*]', "", job_title)
            filename = f"Cover letter - {safe_title} - Cyril Gendarme.pdf"
            export_dir = Path(self._export_directory_var.get())
            export_dir.mkdir(parents=True, exist_ok=True)
            self._data_store.save_export_directory(str(export_dir))
            path = export_dir / filename

            doc = SimpleDocTemplate(
                str(path), pagesize=A4, topMargin=0.5 * inch, bottomMargin=0.5 * inch
            )
            styles = getSampleStyleSheet()

            # Create custom style with Calibri-like font (using Helvetica) at 11pt
            from reportlab.lib.styles import ParagraphStyle

            custom_style = ParagraphStyle(
                "CustomNormal",
                parent=styles["Normal"],
                fontName="Helvetica",
                fontSize=11,
                leading=14,
                spaceBefore=0,
                spaceAfter=0,
            )

            story = []
            import re as regex_module

            for line in text.split("\n"):
                if line == "":
                    story.append(Spacer(1, custom_style.leading))
                    continue

                # Keep user-provided <b>...</b> tags and also bold tech bullet labels.
                formatted_line = self._format_pdf_line_with_bold_tags(line)
                story.append(Paragraph(formatted_line, custom_style))

            doc.build(story)
            self._status_var.set(f"PDF export complete: {path}")
        except ImportError as ie:
            self._status_var.set(f"Missing package: {ie}")
        except Exception as e:
            self._status_var.set(f"Export error: {e}")

    def _get_export_source_text(self) -> str:
        current_output_text = self._output.get("1.0", "end-1c")
        if not self._generated_output_with_tags:
            return current_output_text

        generated_without_tags = self._strip_bold_tags(self._generated_output_with_tags)
        if current_output_text.strip() == generated_without_tags.strip():
            return self._generated_output_with_tags
        return current_output_text

    def _strip_bold_tags(self, text: str) -> str:
        import re

        return re.sub(r"</?b>", "", text, flags=re.IGNORECASE)

    def _add_word_runs_with_bold_tags(
        self, paragraph: object, text: str, default_bold: bool = False
    ) -> None:
        import re

        chunks = re.split(r"(<b>.*?</b>)", text, flags=re.IGNORECASE)
        for chunk in chunks:
            if not chunk:
                continue

            match = re.fullmatch(r"<b>(.*?)</b>", chunk, flags=re.IGNORECASE)
            if match:
                run = paragraph.add_run(match.group(1))
                run.bold = True
                continue

            run = paragraph.add_run(chunk)
            if default_bold:
                run.bold = True

    def _format_pdf_line_with_bold_tags(self, line: str) -> str:
        import re
        from xml.sax.saxutils import escape

        escaped = escape(line)
        escaped = escaped.replace("&lt;b&gt;", "<b>").replace("&lt;/b&gt;", "</b>")
        return re.sub(r"^(\s*-\s*)([^:]+)(:)", r"\1<b>\2</b>\3", escaped)

    def start(self) -> None:
        self._root.mainloop()
